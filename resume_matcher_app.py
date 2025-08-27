import random
import re
import regex
from datetime import datetime
import math
import torch
import os
import matplotlib.pyplot as plt
import matplotlib
import json
import nltk
#from deep_translator import GoogleTranslator
from ftfy import fix_text
import math
from nltk.corpus import stopwords
import img2pdf
import re
import dateparser
from datetime import datetime
#import re
import logging
import math
import argparse
from datetime import datetime
import streamlit as st
import dateparser
import pdfplumber
import docx
import spacy
from langdetect import detect, detect_langs, DetectorFactory
from sentence_transformers import SentenceTransformer, util
from spacy.matcher import PhraseMatcher
import pandas as pd
from transformers import pipeline
import re
import html
import whisper
from transformers import BartTokenizer, BartForConditionalGeneration
import regex  # not 're', so we keep Unicode regex features
import os
import requests
import time
from PIL import Image
from docling.document_converter import DocumentConverter
import pytesseract
MODEL_PATH = 'fine_tuned_bart2F'
CLASSIFIER_DIR = "savedmodel"              # Path to your classifier
TRANSLATION_MODEL = "Helsinki-NLP/opus-mt-fr-en"  # French -> English

import joblib
import os
import nltk
from nltk.corpus import stopwords
from ftfy import fix_text
import re
from transformers import AutoTokenizer, AutoModelForSequenceClassification
import torch
import os
import logging
import unicodedata
import re
import os
import logging
import unicodedata
import re
# 📁 Path to your folder
import logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)
import unicodedata
import re
import torch
import numpy as np
from typing import Tuple


def extract_text_from_file(file_path: str, use_easyocr=False) -> str:
    """
    Extract text from PDF, DOCX/DOC, TXT, or image files.
    Falls back to Docling if standard extraction fails.
    """
    import os
    ext = os.path.splitext(file_path)[1].lower()
    extracted_text = ""

    try:
        if ext == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    extracted_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
        elif ext in [".docx", ".doc"]:
            try:
                from docx import Document
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs.append(cell.text)
                extracted_text = "\n".join(p.strip() for p in paragraphs if p.strip())
            except Exception as e:
                logger.warning(f"docx extraction failed: {e}")
        elif ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    extracted_text = f.read()
            except Exception as e:
                logger.warning(f"txt read failed: {e}")
        elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            try:
                from PIL import Image
                import pytesseract
                img = Image.open(file_path)
                extracted_text = pytesseract.image_to_string(img, lang='fra+eng')
            except Exception as e:
                logger.warning(f"image OCR failed: {e}")
    except Exception as e:
        logger.error(f"Primary extraction failed: {e}")

    # If primary extraction failed or returned empty, try Docling
    if not extracted_text.strip():
        try:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(file_path)
            md = result.document.export_to_markdown()
            if isinstance(md, bytes):
                md = md.decode('utf-8', errors='replace')
            extracted_text = md.strip()
            logger.info("Fallback to Docling extraction succeeded.")
        except Exception as e:
            logger.error(f"Docling extraction failed: {e}")

    # Normalize and deduplicate
    if extracted_text:
        extracted_text = normalize_text(extracted_text)
        extracted_text = deduplicate_sections(extracted_text)

    return extracted_text.strip()

#llama
# recommande : obtenir la clé depuis l'environnement
import os, json, re, time, random, requests
from typing import Tuple, List
import logging

logger = logging.getLogger(__name__)

# Primary and backup models
PRIMARY_MODEL = "llama-3.3-70b-versatile"
BACKUP_MODEL = "llama-3.1-8b-instant"  # smaller model fallback

# Default API keys / endpoints
API_KEYS = [
    os.environ.get("LLM_API_KEY", "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"),
    os.environ.get("LLM_API_KEY_2", "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")  # optional backup key
]
ENDPOINTS = [
    "https://api.groq.com/openai/v1/chat/completions",
    # Add backup endpoints if available
]
def llama_score_resume(
    resume_text: str,
    job: dict,
    model_id: str = PRIMARY_MODEL,
    timeout: int = 30,
    max_retries: int = 3,
    api_keys: List[str] = None
) -> Tuple[float, str]:

    keys_to_use = api_keys or API_KEYS
    endpoints_to_use = ENDPOINTS
    last_err = None

    # Include experience_fields if present
    exp_fields = ", ".join(job.get("experience_fields", []))

    job_prompt = (
        f"Offre d'emploi :\n"
        f"Titre: {job.get('job_title','')}\n"
        f"Description: {job.get('description','')}\n"
        f"Compétences requises: {', '.join(job.get('required_skills',[]))}\n"
        f"Langues requises: {', '.join(job.get('required_languages',[]))}\n"
        f"Diplôme requis: {job.get('required_degree','')}\n"
        f"Expérience (années): {job.get('min_experience_years','')} - {job.get('max_experience_years','')}\n"
        f"Champs d'expérience recherchés: {exp_fields}\n\n"
    )

    for key in keys_to_use:
        for endpoint in endpoints_to_use:
            for attempt in range(1, max_retries + 1):
                try:
                    prompt = (
                        "Vous êtes un évaluateur expert qui note sur 0 à 100 dans quelle mesure un CV correspond à une offre d'emploi.\n"
                        "Retournez uniquement un objet JSON valide EXACTEMENT au format : "
                        '{"score": <entier 0-100>, "explanation": "<courte explication en français>"}\n\n'
                        + job_prompt +
                        f"CV (texte) :\n{resume_text}"
                    )

                    payload = {
                        "model": model_id,
                        "messages": [
                            {"role": "system", "content": "Vous êtes un assistant d'évaluation CV en français."},
                            {"role": "user", "content": prompt}
                        ],
                        "temperature": 0.0,
                        "max_tokens": 256
                    }

                    headers = {
                        "Authorization": f"Bearer {key}",
                        "Content-Type": "application/json"
                    }

                    resp = requests.post(endpoint, headers=headers, json=payload, timeout=timeout)

                    if resp.status_code == 429:
                        wait = (2**attempt) + random.random()
                        logger.warning("429 rate limit. Waiting %.1fs", wait)
                        time.sleep(wait)
                        last_err = f"429 rate limit"
                        continue

                    if 500 <= resp.status_code < 600:
                        wait = (2**attempt) + random.random()
                        logger.warning("Server error %d. Waiting %.1fs", resp.status_code, wait)
                        time.sleep(wait)
                        last_err = f"{resp.status_code} server error"
                        continue

                    resp.raise_for_status()
                    data = resp.json()

                    # Extract content
                    content = ""
                    try:
                        choice = (data.get("choices") or [])[0]
                        msg = choice.get("message") or {}
                        content = msg.get("content") or choice.get("text") or ""
                    except Exception:
                        content = json.dumps(data, ensure_ascii=False)

                    # Parse JSON
                    parsed = None
                    try:
                        parsed = json.loads(content)
                    except:
                        m = re.search(r"(\{(?:.|\n)*\})", content)
                        if m:
                            try: parsed = json.loads(m.group(1))
                            except: parsed = None

                    raw_score = 0
                    expl = content
                    if parsed and isinstance(parsed, dict):
                        raw_score = parsed.get("score", 0)
                        expl = parsed.get("explanation", content)
                    else:
                        m2 = re.search(r"(\d{1,3})", content)
                        if m2:
                            try: raw_score = int(m2.group(1))
                            except: raw_score = 0

                    score_norm = max(0.0, min(1.0, float(raw_score)/100.0))
                    return score_norm, str(expl)

                except Exception as e:
                    last_err = str(e)
                    wait = (2**attempt) + random.random()
                    logger.warning("Attempt %d failed: %s. Waiting %.1fs", attempt, last_err, wait)
                    time.sleep(wait)
                    continue

    # Fallback to backup model if primary failed
    if model_id != BACKUP_MODEL:
        logger.warning("Primary model failed, trying backup model %s", BACKUP_MODEL)
        return llama_score_resume(resume_text, job, model_id=BACKUP_MODEL, timeout=timeout, max_retries=max_retries)

    logger.error("All LLaMA attempts failed. Last error: %s", last_err)
    return 0.0, f"LLM error: {last_err or 'unknown'}"

"""
def llama_score_resume(
    resume_text: str,
    job: dict,
    model_id: str = PRIMARY_MODEL,
    timeout: int = 30,
    max_retries: int = 4,
) -> Tuple[float, str]:

    # Split large CVs into chunks
    max_chunk_len = 4000
    resume_chunks = [resume_text[i:i+max_chunk_len] for i in range(0, len(resume_text), max_chunk_len)]

    # Build job prompt
    job_prompt = (
        f"Offre d'emploi :\n"
        f"Titre: {job.get('job_title','')}\n"
        f"Description: {job.get('description','')}\n"
        f"Compétences requises: {', '.join(job.get('required_skills',[]))}\n"
        f"Langues requises: {', '.join(job.get('required_languages',[]))}\n"
        f"Diplôme requis: {job.get('required_degree','')}\n"
        f"Expérience (années): {job.get('min_experience_years','')} - {job.get('max_experience_years','')}\n\n"
    )

    # Retry logic for multiple keys/endpoints
    last_err = None
    for api_key in API_KEYS:
        for endpoint in ENDPOINTS:
            for attempt in range(1, max_retries+1):
                try:
                    # Send chunks sequentially (you could also aggregate)
                    content_parts = []
                    for chunk in resume_chunks:
                        prompt = (
                            "Vous êtes un évaluateur expert qui note sur 0 à 100 dans quelle mesure un CV correspond à une offre d'emploi.\n"
                            "Retournez uniquement un objet JSON valide EXACTEMENT au format : "
                            '{"score": <entier 0-100>, "explanation": "<courte explication en français>"}\n\n'
                            + job_prompt
                            + f"CV (texte) :\n{chunk}"
                        )

                        payload = {
                            "model": model_id,
                            "messages": [
                                {"role": "system", "content": "Vous êtes un assistant d'évaluation CV en français."},
                                {"role": "user", "content": prompt}
                            ],
                            "temperature": 0.0,
                            "max_tokens": 256
                        }

                        headers = {
                            "Authorization": f"Bearer {api_key}",
                            "Content-Type": "application/json"
                        }

                        resp = requests.post(endpoint, headers=headers, json=payload, timeout=timeout)

                        if resp.status_code == 429:
                            retry_after = resp.headers.get("Retry-After")
                            wait = float(retry_after) if retry_after and retry_after.isdigit() else (2**attempt)+random.random()
                            logger.warning("429 rate limit. Waiting %.1fs", wait)
                            time.sleep(wait)
                            last_err = f"429 rate limit (attempt {attempt})"
                            raise Exception(last_err)

                        if 500 <= resp.status_code < 600:
                            wait = (2**attempt)+random.random()
                            logger.warning("Server error %d. Waiting %.1fs", resp.status_code, wait)
                            time.sleep(wait)
                            last_err = f"{resp.status_code} server error"
                            raise Exception(last_err)

                        resp.raise_for_status()
                        data = resp.json()

                        # Extract OpenAI-like content
                        content = ""
                        try:
                            choices = data.get("choices", [])
                            if choices:
                                choice = choices[0]
                                msg = choice.get("message") or choice.get("delta") or {}
                                content = (msg.get("content") if isinstance(msg, dict) else None) or choice.get("text") or ""
                        except Exception:
                            content = json.dumps(data, ensure_ascii=False)

                        content_parts.append(content.strip())

                    # Combine all chunk responses
                    combined_content = " ".join(content_parts)

                    # Parse JSON safely
                    parsed = None
                    try:
                        parsed = json.loads(combined_content)
                    except Exception:
                        m = re.search(r"(\{(?:.|\n)*\})", combined_content)
                        if m:
                            try: parsed = json.loads(m.group(1))
                            except: parsed = None

                    raw_score = 0
                    expl = combined_content
                    if parsed and isinstance(parsed, dict):
                        raw_score = parsed.get("score", parsed.get("score_value", 0))
                        expl = parsed.get("explanation", parsed.get("comment", combined_content))
                    else:
                        m2 = re.search(r"(\d{1,3})", combined_content)
                        if m2:
                            try: raw_score = int(m2.group(1))
                            except: raw_score = 0

                    score_norm = max(0.0, min(1.0, float(raw_score)/100.0))
                    return score_norm, str(expl)

                except Exception as e:
                    last_err = str(e)
                    wait = (2**attempt)+random.random()
                    logger.warning("Attempt %d failed: %s. Waiting %.1fs", attempt, last_err, wait)
                    time.sleep(wait)
                    continue

    # If all attempts fail, fallback to smaller model if not already
    if model_id != BACKUP_MODEL:
        logger.warning("Primary model failed, trying backup model %s", BACKUP_MODEL)
        return llama_score_resume(resume_text, job, model_id=BACKUP_MODEL, timeout=timeout, max_retries=max_retries)

    logger.error("All LLaMA attempts failed. Last error: %s", last_err)
    return 0.0, f"LLM error: {last_err or 'unknown'}"
"""
"""
DEFAULT_LLM_API_KEY = os.environ.get("LLM_API_KEY", "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")  # remplacer / retirer si tu utilises Config

def llama_score_resume(
    resume_text: str,
    job: dict,
    api_key: str = None,
    model_id: str = "llama-3.3-70b-versatile",
    endpoint: str = "https://api.groq.com/openai/v1/chat/completions",
    timeout: int = 30,
    max_retries: int = 4
) -> Tuple[float, str]:
    
    #Appelle le LLM pour retourner (score_norm_float_0_1, explanation_str).
    #Gère 429 / 5xx / timeout avec retry/backoff et parse la réponse de façon tolérante.
    

    if api_key is None:
        api_key = DEFAULT_LLM_API_KEY

    # Prompt clair demandant un JSON strict — continue de le garder
    prompt = (
        "Vous êtes un évaluateur expert qui note sur 0 à 100 dans quelle mesure un CV correspond à une offre d'emploi.\n"
        "Retournez uniquement un objet JSON valide EXACTEMENT au format suivant : "
        '{"score": <entier 0-100>, "explanation": "<courte explication en français>"}\n\n'
        "Offre d'emploi :\n"
        f"Titre: {job.get('job_title','')}\n"
        f"Description: {job.get('description','')}\n"
        f"Compétences requises: {', '.join(job.get('required_skills',[]))}\n"
        f"Langues requises: {', '.join(job.get('required_languages',[]))}\n"
        f"Diplôme requis: {job.get('required_degree','')}\n"
        f"Expérience (années): {job.get('min_experience_years','')} - {job.get('max_experience_years','')}\n\n"
        "CV (texte) :\n" + resume_text
    )

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model_id,
        "messages": [
            {"role": "system", "content": "Vous êtes un assistant d'évaluation CV en français."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.0,
        "max_tokens": 256
    }

    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            resp = requests.post(endpoint, headers=headers, json=payload, timeout=timeout)

            # 429: rate limit -> respectez Retry-After si présent
            if resp.status_code == 429:
                retry_after = resp.headers.get("Retry-After")
                wait = float(retry_after) if retry_after and retry_after.isdigit() else (2 ** attempt) + random.random()
                logger.warning("LLM 429 rate limit on attempt %d. Waiting %.1fs", attempt, wait)
                time.sleep(wait)
                last_err = f"429 rate limit (attempt {attempt})"
                continue

            # 5xx: serveur malade -> retry with backoff
            if 500 <= resp.status_code < 600:
                wait = (2 ** attempt) + random.random()
                logger.warning("LLM server error %d on attempt %d, waiting %.1fs", resp.status_code, attempt, wait)
                time.sleep(wait)
                last_err = f"{resp.status_code} server error"
                continue

            resp.raise_for_status()
            data = resp.json()

            # essaye d'extraire le texte selon format OpenAI-like
            content = None
            if isinstance(data, dict):
                choices = data.get("choices") or []
                if choices and isinstance(choices, list):
                    # support safer indexing
                    choice = choices[0]
                    # OpenAI-like: choice['message']['content']
                    if isinstance(choice, dict):
                        msg = choice.get("message") or choice.get("delta") or {}
                        content = (msg.get("content") if isinstance(msg, dict) else None) or choice.get("text") or ""
            if content is None:
                # dernier recours : stringify whole response
                content = json.dumps(data, ensure_ascii=False)

            content = content.strip()
        except requests.exceptions.Timeout as e:
            last_err = f"timeout: {e}"
            wait = (2 ** attempt) + random.random()
            logger.warning("Timeout calling LLM (attempt %d). Waiting %.1fs", attempt, wait)
            time.sleep(wait)
            continue
        except Exception as e:
            last_err = str(e)
            wait = (2 ** attempt) + random.random()
            logger.exception("Unexpected error calling LLM (attempt %d): %s. Waiting %.1fs", attempt, e, wait)
            time.sleep(wait)
            continue

        # ---------- parse content ----------
        # 1) try strict JSON first
        parsed = None
        try:
            parsed = json.loads(content)
        except Exception:
            # 2) try to extract JSON substring { ... } with regex
            m = re.search(r"(\{(?:.|\n)*\})", content)
            if m:
                try:
                    parsed = json.loads(m.group(1))
                except Exception:
                    parsed = None

        # 3) fallback: try to extract first integer found
        raw_score = 0
        expl = content
        if parsed and isinstance(parsed, dict):
            raw_score = parsed.get("score", parsed.get("score_value", 0))
            expl = parsed.get("explanation", parsed.get("comment", content))
        else:
            m2 = re.search(r"(\d{1,3})", content)
            if m2:
                try:
                    raw_score = int(m2.group(1))
                except Exception:
                    raw_score = 0

        # normalize to 0.0 - 1.0
        try:
            score_norm = float(raw_score) / 100.0
        except Exception:
            try:
                score_norm = float(raw_score)
            except Exception:
                score_norm = 0.0

        score_norm = max(0.0, min(1.0, score_norm))

        return score_norm, str(expl)

    # si on sort de la boucle : erreur persistante
    logger.error("LLM scoring failed after %d attempts. Last error: %s", max_retries, last_err)
    return 0.0, f"LLM error: {last_err or 'unknown'}"
"""
import re

# --- Normalization helper ---
def normalize_degree(text: str) -> str:
    if not text:
        return ""
    t = text.lower().strip()

    mapping = {
        "phd": ["phd", "ph.d", "doctorat", "docteur", "dr", "thèse", "these"],
        "master": [
            "master", "mastère", "mastere", "maîtrise", "maitrise", "msc", "mba", "m2", "m1",
            "master spécialisé", "mastère spécialisé", "master specialise", "mastère specialise",
            "master professionnel", "master pro", "mastère professionnel", "mastère pro",
            "master de recherche"
        ],
        "engineer": [
            "ingénieur", "ingenieur", "ing", "ingénieur d'état", "ingenieur d'etat",
            "diplôme d’ingénieur", "diplome d'ingenieur", "cycle d’ingénieur", "cycle d'ingenieur"
        ],
        "bachelor": ["bachelor", "licence", "licence pro", "licence professionnelle", "b.sc", "b.s", "licenciatura"],
        "baccalaureat": ["baccalauréat", "baccalaureat", "bac"],
        "deug": ["deug", "deust", "dut", "bts"]
    }

    for canon, variants in mapping.items():
        for v in variants:
            if v in t:
                return canon
    return ""

# --- Education extraction ---
def extract_education(text: str):
    education_entries = []

    # normalize text: unify dashes and remove newlines
    text = text.replace("–", "-").replace("—", "-").replace("\n", " ")

    # flexible regex to catch both "YYYY - YYYY", "YYYY YYYY", or "YYYY en cours"
    edu_pattern = re.compile(r"(\d{4})\s*(?:-\s*| )(\d{4}|en cours)?\s*(.+?)(?=(\d{4}\s*(?:-| )|\Z))", re.IGNORECASE)

    for m in edu_pattern.finditer(text):
        year1, year2, chunk, _ = m.groups()
        deg_norm = normalize_degree(chunk)
        education_entries.append({
            "years": [year1, year2] if year2 else [year1],
            "degree_raw": chunk.strip(),
            "degree_canon": deg_norm
        })

    return education_entries

# --- Scoring ---
def score_education(resume_text, job_required_degree, preferred_fields):
    hits = extract_education(resume_text)
    best_score = 0.0
    best_hit = None

    for h in hits:
        score = 0.0
        # degree match
        if h["degree_canon"] == job_required_degree:
            score += 0.6
        # field match
        for f in preferred_fields:
            if f.lower() in h["degree_raw"].lower():
                score += 0.4
                break
        if score > best_score:
            best_score = score
            best_hit = h

    return {"score": best_score, "best": best_hit, "merged": hits}

from typing import List, Union, Dict
import re, unicodedata, html
from datetime import datetime
import dateparser
from sentence_transformers import SentenceTransformer, util
from typing import Union, List, Dict
import re, html, unicodedata
from datetime import datetime
import dateparser
from sentence_transformers import SentenceTransformer, util

# Load multilingual Sentence-BERT model once
model = SentenceTransformer("sentence-transformers/distiluse-base-multilingual-cased-v2")

def full_experience(raw_text: str,
                    min_years: float,
                    max_years: float,
                    job_titles: Union[str, List[str]]) -> Dict:

    # ------------------- Helpers -------------------
    def normalize_text(s: str) -> str:
        if not s:
            return ""
        s = html.unescape(s)
        s = unicodedata.normalize("NFKD", s)
        s = "".join(c for c in s if not unicodedata.combining(c))
        s = s.lower()
        s = re.sub(r"[^\w\s]", " ", s)
        s = re.sub(r"\s+", " ", s).strip()
        return s

    def extract_experience_section(text: str) -> str:
        """Try to isolate the experience section, fallback to whole text"""
        norm_text = normalize_text(text)
        pattern = r"(exp[ée]riences?|parcours professionnel|stages?|projets professionnels|work experience|professional experience)(.*?)(?=\n[A-Z]{2,}|\bformation\b|\beducation\b|$)"
        m = re.search(pattern, norm_text, flags=re.DOTALL)
        if m:
            start, end = m.span()
            return text[start:end]
        return text   # fallback: whole CV

    def parse_experiences(section: str) -> List[Dict]:
        """Parse the experience section into structured entries"""
        MONTHS_PATTERN = r"(Janv|F[ée]vr|Mars|Avr|Mai|Juin|Juil|Ao[uû]t|Sept|Oct|Nov|D[ée]c)[a-zé]*"
        entries = []
        lines = [l.strip() for l in section.split("\n") if l.strip()]
        current_entry = None

        for line in lines:
            line_norm = normalize_text(line)

            # Case 1: date ranges like "Mars 2020 – Juin 2023"
            date_match = re.search(
              rf"({MONTHS_PATTERN}\s+\d{{4}})\s*(?:[-–]|\s+)\s*({MONTHS_PATTERN}\s+\d{{4}}|Present|Aujourd'hui)",
              line, flags=re.IGNORECASE
            )

            if date_match:
                if current_entry:
                    entries.append(current_entry)
                start_str, end_str = date_match.group(1), date_match.group(2)
                start_date = dateparser.parse(start_str, languages=["fr", "en"])
                end_date = dateparser.parse(end_str, languages=["fr", "en"]) or datetime.now()
                title_company_part = line[date_match.end():].strip(" –-")
                current_entry = {
                    "title": title_company_part or "Titre non spécifié",
                    "company": title_company_part.split("–")[-1].strip() if "–" in title_company_part else "",
                    "start_date": start_date,
                    "end_date": end_date,
                    "description": "",
                    "duration_years": round((end_date - start_date).days / 365.0, 2) if start_date and end_date else 0
                }
                continue

            # Case 2: explicit duration like "3 ans" or "6 mois"
            dur_match = re.search(r"(?:\b|[-•])\s*(\d{1,2})\s*(ans|mois)\b", line_norm)
            if dur_match and current_entry:
                number, unit = int(dur_match.group(1)), dur_match.group(2)
                years = number if unit.lower().startswith("an") else round(number / 12.0, 2)
                current_entry["duration_years"] = (current_entry.get("duration_years") or 0) + years
                continue

            # Case 3: continuation text
            if current_entry:
                current_entry["description"] = (current_entry.get("description","") + " " + line).strip()

        if current_entry:
            entries.append(current_entry)
        return entries

    def is_title_matching(title: str, job_titles: List[str]) -> bool:
        """Check if a title matches any job title using normalized substring"""
        title_norm = normalize_text(title)
        for jt in job_titles:
            jt_norm = normalize_text(jt)
            if re.search(rf"\b{re.escape(jt_norm)}\b", title_norm):
                return True
        return False

    def score_total_experience(entries: List[Dict], min_years: float, max_years: float, job_titles: List[str]):
        if isinstance(job_titles, str):
            job_titles = [job_titles]

        # Normalize job titles
        job_titles_norm = [normalize_text(j) for j in job_titles]

        # Compute relevant years using title match + optional semantic similarity
        relevant_years = 0.0
        sim_scores = []

        # Encode job titles once
        job_embs = model.encode(job_titles_norm, convert_to_tensor=True)

        for e in entries:
            dur = e.get("duration_years") or 0
            if dur == 0:
                continue

            title_norm = normalize_text(e.get("title",""))
            # Check exact/substring match
            matched = any(jt in title_norm for jt in job_titles_norm)

            # Compute semantic similarity if no exact match
            if not matched:
                title_emb = model.encode(title_norm, convert_to_tensor=True)
                sim = util.cos_sim(title_emb, job_embs).max().item()
                if sim > 0.7:  # threshold, can tweak
                    matched = True

            if matched:
                relevant_years += dur

        # Normalize experience relative to max_years
        exp_score = min(relevant_years / max_years, 1.0)
        return relevant_years, exp_score

    # ------------------- Main -------------------
    section = extract_experience_section(raw_text)
    entries = parse_experiences(section)

    total_years, exp_score = score_total_experience(entries, min_years, max_years, job_titles)

    return {
        "entries": entries,
        "total_years": round(total_years, 1),
        "score": round(exp_score, 2)
    }


# ─── Classifier ───────────────────────────────────────────────────────────────

# ─── Load once ─────────────────────────────────────────────────────────────────
# ─── Classifier with robust translation ────────────────────────────────────────
import streamlit as st
import torch
import joblib
from transformers import AutoTokenizer, AutoModelForSequenceClassification, AutoModelForSeq2SeqLM
import os
import json
from groq import Groq

# ---------------------- Groq Client ----------------------
client = Groq(api_key="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")

# ---------------------- Classifier CV français ----------------------
def classify_cv(french_text: str):
    """
    Classifie un CV français directement avec LLaMA sur Groq.
    Retourne : label prédit, score de confiance, dictionnaire de probabilités
    """
    prompt = f"""
Vous êtes un assistant de recrutement. Classez le CV suivant en français dans une catégorie de poste.
Répondez UNIQUEMENT au format JSON avec les clés suivantes :
{{ "category": "<nom de la catégorie de poste>", "confidence": <score entre 0 et 1> }}

CV :
{french_text}
    """

    response = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile"
    )

    output_text = response.choices[0].message.content.strip()

    # Parsing JSON sécurisé
    try:
        result = json.loads(output_text)
        pred_label = result.get("category", "Unknown")
        top_prob = float(result.get("confidence", 0.0))
    except Exception:
        pred_label = "Unknown"
        top_prob = 0.0

    prob_dict = {pred_label: top_prob}
    return pred_label, top_prob, prob_dict

# ---------------- Translation ----------------
"""@st.cache_resource
def load_translation_model():
    tok = AutoTokenizer.from_pretrained(TRANSLATION_MODEL, force_download=True)
    model = AutoModelForSeq2SeqLM.from_pretrained(TRANSLATION_MODEL, force_download=True)
    model.eval()

    return tok, model

def translate_resume(french_text: str) -> str:
    
    tok, model = load_translation_model()

    # Tokenize without truncation
    tokens = tok(french_text, return_tensors="pt", truncation=False)["input_ids"][0]

    # Chunk into <=512 tokens
    max_len = model.config.max_position_embeddings  # usually 512
    chunks = [tokens[i:i+max_len] for i in range(0, len(tokens), max_len)]

    translations = []
    for chunk in chunks:
        inputs = {"input_ids": chunk.unsqueeze(0)}
        outputs = model.generate(**inputs, max_length=max_len)
        translated = tok.decode(outputs[0], skip_special_tokens=True)
        translations.append(translated)

    return " ".join(translations)

# ---------------- Classifier ----------------
@st.cache_resource
def load_classifier():
    clf_tok = AutoTokenizer.from_pretrained(CLASSIFIER_DIR)
    clf_model = AutoModelForSequenceClassification.from_pretrained(
        CLASSIFIER_DIR,
        device_map="auto",
        low_cpu_mem_usage=True
    )
    clf_model.eval()
    clf_le = joblib.load(f"{CLASSIFIER_DIR}/label_encoder.pkl")
    return clf_tok, clf_model, clf_le

def classify_cv(english_text: str):

    clf_tokenizer, clf_model, clf_label_encoder = load_classifier()

    inputs = clf_tokenizer(
        english_text,
        return_tensors="pt",
        truncation=True,
        padding="max_length",
        max_length=512
    )

    with torch.no_grad():
        logits = clf_model(**inputs).logits.squeeze(0)

    probs = torch.softmax(logits, dim=0).cpu().numpy()
    classes = list(clf_label_encoder.classes_)

    top_idx = int(probs.argmax())
    pred_label = classes[top_idx]
    top_prob = float(probs[top_idx])
    prob_dict = {c: float(p) for c, p in zip(classes, probs)}

    return pred_label, top_prob, prob_dict

# ---------------- Wrapper: Translate + Classify ----------------

def classify_cv_with_translation(text: str, is_french: bool = True):
    if is_french:
        text = translate_resume(text)  # translate to English first
        print("Translated resume preview:\n", text[:500])

    # Correct unpacking
    pred_label, top_prob, prob_dict = classify_cv(text)
    return pred_label, top_prob, prob_dict

"""#hna
"""
@st.cache_resource
def load_classifier():
    clf_tok   = AutoTokenizer.from_pretrained(CLASSIFIER_DIR)
    # 👉 Ajout d'options pour réduire l'utilisation mémoire
    clf_model = AutoModelForSequenceClassification.from_pretrained(
        CLASSIFIER_DIR,
        device_map="auto",          # Choisit automatiquement CPU ou GPU dispo
        low_cpu_mem_usage=True      # Réduit la charge mémoire au chargement
    )
    clf_model.eval()
    # Keep model on CPU by default. If you want GPU later:
    # device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    # clf_model.to(device)
    clf_le    = joblib.load(f"{CLASSIFIER_DIR}/label_encoder.pkl")
    return clf_tok, clf_model, clf_le

def get_classifier():
    return load_classifier()

def classify_cv(text: str):
   
    clf_tokenizer, clf_model, clf_label_encoder = get_classifier()

    # Tokenize
    inputs = clf_tokenizer(
        text,
        return_tensors="pt",
        truncation=True,
        padding="max_length",
        max_length=512
    )

    # If you moved model to GPU, also send inputs to device:
    # device = next(clf_model.parameters()).device
    # inputs = {k:v.to(device) for k,v in inputs.items()}

    with torch.no_grad():
        logits = clf_model(**inputs).logits.squeeze(0)   # shape: (num_classes,)

    probs = torch.softmax(logits, dim=0).cpu().numpy()   # numpy array
    classes = list(clf_label_encoder.classes_)          # same order as probs

    top_idx = int(probs.argmax())
    pred_lbl = classes[top_idx]
    top_prob = float(probs[top_idx])

    prob_dict = {c: float(p) for c, p in zip(classes, probs)}
    return pred_lbl, top_prob, prob_dict
"""
#lhna
#-------------------AUDIO------------------:
# Load the model once
whisper_model = None  # Do not load yet

def transcribe_audio(audio_path):
    """
    Convert audio file to text using Whisper
    """
    global whisper_model
    try:
        if whisper_model is None:
            import whisper
            whisper_model = whisper.load_model("base")
        result = whisper_model.transcribe(audio_path)
        return result["text"]
    except Exception as e:
        return f"[ERROR] Transcription failed: {e}"

# ─── Summarizer ───────────────────────────────────────────────────────────────


_mojibake_marker = re.compile(r'(Ã|Â)[\x80-\xBF]')

def fix_mojibake_text(s: str) -> str:
    """Essaye ftfy puis un fallback latin1->utf8 si on détecte des marqueurs de mojibake."""
    if not s:
        return s
    # 1) ftfy (meilleur choix en général)
    try:
        fixed = fix_text(s)
        if fixed and fixed != s:
            return fixed
    except Exception:
        pass

    # 2) heuristique: si on voit 'Ã' ou 'Â' suivi d'un octet typique -> tenter re-decode
    if _mojibake_marker.search(s):
        try:
            candidate = s.encode('latin-1').decode('utf-8')
            # sanity: préférer candidate si elle a plus de caractères non-ascii utiles
            if sum(1 for ch in candidate if ord(ch) > 127) >= sum(1 for ch in s if ord(ch) > 127):
                return candidate
        except Exception:
            pass

    # 3) retour du texte inchangé
    return s

# make sure you have the stopwords ready
nltk.download('stopwords')
# build your french+english stop-word set once, globally
stop_words = set(stopwords.words('french') + stopwords.words('english'))
@st.cache_resource
def load_summarizer():
    summ_tok   = BartTokenizer.from_pretrained(MODEL_PATH)
    summ_model = BartForConditionalGeneration.from_pretrained(MODEL_PATH)
    summ_model.eval()
    return summ_model, summ_tok

summ_model, summ_tokenizer = load_summarizer()

#text-extract:

"""
def nettoyer_texte(texte):
    # 1) Remove <!-- image --> tag
    texte = regex.sub(r'<!-- image -->', ' ', texte)

    # 2) Keep only:
    #    - Letters (\p{L}) → all alphabets, including accents
    #    - Digits (\p{N})
    #    - Spaces, colons, commas, dots, parentheses, dashes, slashes, newlines
    texte = regex.sub(r"[^\p{L}\p{N}\s:,\.\(\)\-–/]", " ", texte)

    # 3) Collapse spaces/tabs (but keep newlines)
    texte = regex.sub(r"[ \t]+", " ", texte)

    # 4) Trim each line and remove empties
    lines = [ln.strip() for ln in texte.split("\n")]
    texte = "\n".join([ln for ln in lines if ln])

    # 5) Remove stop words (case-insensitive)
    mots = texte.split()
    texte = " ".join(w for w in mots if w.lower() not in stop_words)
    print("nettoyage:",texte)
    return texte

def extract_text_with_docling(file_path):
    try:
        converter = DocumentConverter()
        result = converter.convert(file_path)
        md = result.document.export_to_markdown()
        # if bytes, decode; if str, keep
        if isinstance(md, bytes):
            md = md.decode('utf-8', errors='replace')
        return md.strip()
    except Exception as e:
        logging.error(f"Erreur Docling : {e}")
        return None
def nettoyer_texte(txt):
    allowed_chars = "A-Za-zÀ-ÖØ-öø-ÿ0-9\s\.,;:!?\(\)\-\+\/#@&_%"  
    txt = re.sub(f"[^{allowed_chars}]", " ", txt)
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt

# Fonction pour convertir une image en PDF
def convert_image_to_pdf(image_path):
    pdf_path = os.path.splitext(image_path)[0] + ".pdf"
    with open(pdf_path, "wb") as pdf_file:
        pdf_file.write(img2pdf.convert(image_path))
    return pdf_path


def detect_language(text):
    try:
        return detect(text)
    except Exception as e:
        logging.error(f"Erreur de détection de langue : {e}")
        return 'fr'  # Par défaut, suppose que le texte est en français
"""
def parse_date_str(s: str):
    if not s:
        return None
    return dateparser.parse(s, languages=['fr', 'en'])

def chunk_text(text: str, max_tokens: int):
    words = text.split()
    for i in range(0,len(words),max_tokens): yield ' '.join(words[i:i+max_tokens])

def parse_date_str(s:str): return dateparser.parse(s) if s else None


def translate_text(text, source_lang='en', target_lang='fr'):
    
   # Traduit un texte donné de la langue source à la langue cible.
    
    try:
        translator = GoogleTranslator(source=source_lang, target=target_lang)
        return translator.translate(text)
    except Exception as e:
        logging.error(f"Erreur de traduction : {e}")
        return text  # Retourne le texte original en cas d'erreur
        
def translate_text_chunked(text, source_lang='en', target_lang='fr', max_words=100):
    translated_chunks = []
    for chunk in chunk_text(text, max_words):
        translated_chunk = translate_text(chunk, source_lang, target_lang)
        translated_chunks.append(translated_chunk)
    return " ".join(translated_chunks)


# ----------------- Logging -----------------
logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s:%(name)s: %(message)s')
logger = logging.getLogger("ResumeExtractor")

# ----------------- Normalization -----------------
def normalize_text(text):
    # Keep newlines, collapse multiple spaces per line
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join([line for line in lines if line.strip()])


def deduplicate_sections(text: str) -> str:
    """Remove duplicate sections based on headings."""
    seen = set()
    lines = []
    for line in text.splitlines():
        if ":" in line:
            section = line.split(":")[0].strip().lower()
            if section in seen:
                continue
            seen.add(section)
        lines.append(line)
    return "\n".join(lines)

# ----------------- PDF Extraction -----------------
import os, re, unicodedata, logging, spacy, subprocess, sys
from langdetect import detect

log = logging.getLogger("ResumeExtractor")
logging.basicConfig(level=logging.INFO)

# ----------------- spaCy loader -----------------
def load_spacy_model():
    try:
        return spacy.load("fr_core_news_sm")
    except OSError:
        log.warning("spaCy French model not found. Installing...")
        subprocess.run([sys.executable, "-m", "spacy", "download", "fr_core_news_sm"], check=True)
        return spacy.load("fr_core_news_sm")
    except Exception as e:
        log.error(f"Failed to load spaCy model: {e}", exc_info=True)
        return None

# Load model ONCE at startup
nlp = load_spacy_model()

# ----------------- Normalization -----------------
def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = re.sub(r"\s+", " ", text)
    return text.strip()

# ----------------- Main processing -----------------
def process_file(file_path):
    # Use your existing extract_text_from_file here
    extracted_text = extract_text_from_file(file_path)

    # Language detection
    if extracted_text:
        lang = detect(extracted_text)
        if lang == "en":
            extracted_text = translate_text(extracted_text, source_lang="en", target_lang="fr")

    # spaCy processing
    if nlp:
        doc = nlp(extracted_text)
        # You can now use doc.ents, doc.sents, etc.
    else:
        log.warning("Skipping NLP — spaCy not loaded")

    return extracted_text


    # Return the longest
    return max(candidates, key=len, default="")
from docx import Document

def extract_text_from_docx(path):
    doc = Document(path)
    text = []

    for p in doc.paragraphs:
        text.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)

    return "\n".join(t.strip() for t in text if t.strip())
def nettoyer_texte(txt):
    allowed_chars = r"A-Za-zÀ-ÖØ-öø-ÿ0-9\s\.,;:!?\(\)\-\+\/#@&_%"  
    txt = re.sub(rf"[^{allowed_chars}]", " ", txt)
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt
"""
# ----------------- Unified extractor -----------------
def extract_text_from_file(file_path: str, use_easyocr=False) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    extracted_text = ""

    if ext == ".pdf":
        extracted_text = extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        try:
            import docx
            doc = docx.Document(file_path)
            extracted_text = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            logger.warning(f"docx extraction failed: {e}")
    elif ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                extracted_text = f.read()
        except Exception as e:
            logger.warning(f"txt read failed: {e}")
    elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
        try:
            from PIL import Image
            import pytesseract
            img = Image.open(file_path)
            extracted_text = pytesseract.image_to_string(img, lang='fra+eng')
        except Exception as e:
            logger.warning(f"image OCR failed: {e}")

    # Normalize and deduplicate
    if extracted_text:
        extracted_text = normalize_text(extracted_text)
        extracted_text = deduplicate_sections(extracted_text)

    return extracted_text.strip()
"""
# ----------------- Optional: language detection + translation -----------------
def detect_language(text: str) -> str:
    try:
        return detect(text)
    except:
        return "unknown"

def extract_and_translate(file_path: str, use_easyocr=False, target_lang='fr') -> str:
    """Extract text and translate if English."""
    text = extract_text_from_file(file_path)
    if detect_language(text) == 'en':
        try:
            from googletrans import Translator
            translator = Translator()
            text = translator.translate(text, src='en', dest=target_lang).text
        except Exception as e:
            logger.warning(f"Translation failed: {e}")
    return text

def generate_summary(txt):
    inputs = summ_tokenizer(
        txt,
        return_tensors="pt",
        max_length=1056,
        truncation=True
    )
    summary_ids = summ_model.generate(
        **inputs,
        max_length=556,
        min_length=30,
        num_beams=4,
        early_stopping=True
    )
    raw = summ_tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    repaired = fix_mojibake_text(raw)
    # debug utile (enlève ou commente en production)
    logger.debug("Summary repr: %s", repr(raw)[:300])
    if repaired != raw:
        logger.info("Summary was repaired for mojibake.")
        logger.debug("Repaired repr: %s", repr(repaired)[:300])
    return repaired

import re

import re
"""
def clean_summary(summary_text):
    if not summary_text:
        return ""

    seen_sections = set()
    cleaned_lines = []

    # Regex to detect age patterns in French and English
    age_pattern = re.compile(r'\b\d{1,3}\s*(ans|years\s*old|years)\b', re.IGNORECASE)

    for line in summary_text.splitlines():
        # Remove age mentions but keep the rest of the line
        line = age_pattern.sub("", line).strip()
        if not line:  # skip empty lines after removal
            continue

        # Skip duplicate sections
        if ":" in line:
            section = line.split(":")[0].strip().lower()
            if section in seen_sections:
                continue
            seen_sections.add(section)

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines) if cleaned_lines else ""
"""
import re
import re

def clean_summary(summary_text: str) -> str:
    """
    Nettoie un résumé :
    - supprime les mentions d'âge,
    - supprime les doublons de sections,
    - supprime certaines phrases clés même si elles apparaissent après ':'.
    """
    if not summary_text:
        return ""

    seen_sections = set()
    cleaned_lines = []

    # Regex pour détecter âges en français ou anglais
    age_pattern = re.compile(r'\b\d{1,3}\s*(ans|years\s*old|years)\b', re.IGNORECASE)

    # Phrases à retirer
    phrases_to_remove = [
        "Informations personnelles:",
        "en une phrase concise, en mettant uniquement en avant leur objectif et impact :",
        "avec leur spécialité, établissement et année :",
        "mentionnant le CV sous forme de mots-clés, en évitant les répétitions :",
        "Voici les expériences professionnelles mentionnées dans le CV :",
        "parlées selon le CV sous forme de mots-clés, en évitant les répétitions :"
    ]

    for line in summary_text.splitlines():
        line = line.strip()
        if not line:
            continue

        # Supprimer phrases spécifiques après ':' ou au début
        for phrase in phrases_to_remove:
            # Cherche le phrase à partir de n'importe où dans la ligne, insensible à la casse
            pattern = re.compile(re.escape(phrase), re.IGNORECASE)
            line = pattern.sub("", line).strip()

        # Supprimer mentions d'âge
        line = age_pattern.sub("", line).strip()
        if not line:
            continue

        # Éviter doublons de sections (basé sur le texte avant le premier ':')
        if ":" in line:
            section = line.split(":")[0].strip().lower()
            if section in seen_sections:
                continue
            seen_sections.add(section)

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines) if cleaned_lines else ""


# --------------------- CONFIGURATION ---------------------
class Config:
    MODEL_NAME = "sentence-transformers/distiluse-base-multilingual-cased-v2"

    WEIGHTS = {
       'semantic': 0.25,
       'skills': 0.35,
       'experience': 0.2,
       'education': 0.15,
       'language': 0.05   # small but still counts
    }
    DEGREE_HIERARCHY = {
       'phd': 4, 'doctorat': 4, 'doctorate': 4,
       'master': 3, 'msc': 3, 'maîtrise': 3, 'bac+5': 3, 'mba': 3,
       'engineer': 3, 'ingénieur': 3, 'diplôme d’ingénieur': 3,
       'bachelor': 2, 'licence': 2, 'b.s.': 2, 'bsc': 2,
       'associate': 1, 'bts': 1, 'dut': 1, 'deug': 1, 'deust': 1, 'bac+2': 1
}

    LANG_MAP = {'french':'fr','français':'fr','francais':'fr','english':'en','anglais':'en','arabic':'ar','arabe':'ar','german':'de','allemand':'de',
                'spanish':'es','espagnol':'es','italian':'it','chinese':'zh','mandarin':'zh','japanese':'ja'}
    MODEL_NAME = 'sentence-transformers/distiluse-base-multilingual-cased-v2'
    CHUNK_TOKENS = 512

logging.basicConfig(level=logging.INFO,format='%(asctime)s %(levelname)s:%(name)s: %(message)s')
logger = logging.getLogger('ResumeMatcher')

# Fonction pour extraire le texte d'un fichier à l'aide de Docling


# Assurez-vous que les stopwords sont téléchargés
#nltk.download('stopwords')

# Définir les stop words en français et en anglais
#stop_words = set(stopwords.words('french') + stopwords.words('english'))
#import re

# Example stop words (expand as needed)
# Assurez-vous que les stopwords sont téléchargés

# Définir les stop words en français et en anglais


# Fonction pour traiter les fichiers




# Fonction pour extraire et traiter un fichier ZIP

# ---------- SCORING COMPONENTS ----------
def score_language(text: str, required: set):
    """
    Language scoring:
    - Only uses declared languages in resume (no detected fallback).
    - Score is proportion of required languages declared.
    - If no required languages specified, returns 1.0.
    """
    declared = set()
    # Look for a 'Languages' or 'Langues' section
    m = re.search(r"(?:langues?|languages?)\s*[:\-–]?\s*(.*?)(?:\n{2}|$)",
                  text, re.IGNORECASE | re.DOTALL)
    if m:
        # Extract tokens and map to codes
        for tok in re.findall(r"\b(\w+)\b", m.group(1)):
            code = Config.LANG_MAP.get(tok.lower())
            if code:
                declared.add(code)

    # Compute score strictly from declared
    if not required:
        score = 1.0
    else:
        score = len(declared & required) / len(required) if declared else 0.0

    return round(score, 4), sorted(declared)


def score_skills(text: str, required_skills: list):
    nlp = spacy.blank('fr')
    matcher = PhraseMatcher(nlp.vocab, attr='LOWER')
    # Create patterns only for required skills
    patterns = [nlp.make_doc(skill) for skill in required_skills]
    matcher.add('SKILLS', patterns)
    doc = nlp(text.lower())
    matches = matcher(doc)
    found_skills = set()
    for _, start, end in matches:
        found_skills.add(doc[start:end].text)
    # Score = number of matched required skills / total required skills
    score = len(found_skills) / len(required_skills) if required_skills else 1.0
    return round(score, 4), list(found_skills)

# add imports near top
import html
import unicodedata

def normalize_text_ascii(s: str) -> str:
    """Remove diacritics, lower, collapse spaces."""
    if not s:
        return ""
    # 1) NFKD normalization then strip combining diacritics
    nk = unicodedata.normalize("NFKD", s)
    ascii_only = "".join(ch for ch in nk if not unicodedata.combining(ch))
    return re.sub(r"\s+", " ", ascii_only).strip().lower()

def keyword_match(candidate: str, req_domain: str) -> bool:
    """Simple robust keyword test: normalize both, require all words in req_domain to appear in candidate."""
    if not candidate or not req_domain:
        return False
    cand_norm = normalize_text_ascii(candidate)
    req_norm = normalize_text_ascii(req_domain)
    req_tokens = [t for t in re.split(r'\W+', req_norm) if t]
    return all(tok in cand_norm.split() for tok in req_tokens)

#test

#test



# Patch: improved experience extraction and scoring utilities
# Replace the old extract_experience, score_experience, and classify_experience_domain
# with the functions below. They are designed to be drop-in compatible with the
# rest of your pipeline with minimal changes (see wiring notes at the end).

"""
def score_semantic(text,job_txt,sbert):
    best=0.0
    for chunk in chunk_text(text,Config.CHUNK_TOKENS):
        sim=util.cos_sim(sbert.encode(chunk,convert_to_tensor=True),
                         sbert.encode(job_txt,convert_to_tensor=True)).item()
        best=max(best,sim)
    return round(best,4)
"""
import re
from functools import lru_cache
from sentence_transformers import util

# --- Config par défaut (si Config existe déjà, ajoute seulement les attributs)
# --- Config par défaut (si Config existe déjà, ajoute seulement les attributs) ---
if 'Config' not in globals():
    class Config:
        CHUNK_SIZE = 120
        CHUNK_OVERLAP = 20
else:
    # ne pas redéfinir la classe; définir seulement les attributs manquants
    if not hasattr(Config, "CHUNK_SIZE"):
        Config.CHUNK_SIZE = 120
    if not hasattr(Config, "CHUNK_OVERLAP"):
        Config.CHUNK_OVERLAP = 20


# simple cache mémoire pour embeddings de job (clé légère : id ou tuple(title,company))
_job_emb_cache = {}

def build_job_text(job: dict) -> str:
    """Concatène automatiquement les champs importants du job en un seul texte pour l'embedding."""
    parts = []
    if job.get("title"):
        parts.append(f"Titre: {job.get('title')}")
    if job.get("description"):
        parts.append(f"Description: {job.get('description')}")
    if job.get("required_skills"):
        parts.append("Compétences requises: " + ", ".join(job.get("required_skills")))
    if job.get("preferred_skills"):
        parts.append("Compétences souhaitées: " + ", ".join(job.get("preferred_skills")))
    if job.get("required_degree"):
        parts.append("Diplôme requis: " + str(job.get("required_degree")))
    if job.get("min_experience_years"):
        parts.append("Expérience requise: " + str(job.get("min_experience_years")) + " ans")
    # fallback: any other textual metadata
    if job.get("meta"):  # si tu utilises un champ meta pour infos additionnelles
        parts.append(str(job.get("meta")))
    return "\n".join([p for p in parts if p])

def semantic_chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> list:
    """
    Chunking simple mais sémantique :
    - split par paragraphes puis par phrases (séparateurs .?!)
    - regroupe des phrases pour former des chunks d'environ `chunk_size` mots
    - applique overlap en mots entre chunks
    """
    if chunk_size is None:
        chunk_size = getattr(Config, "CHUNK_SIZE", 120)
    if overlap is None:
        overlap = getattr(Config, "CHUNK_OVERLAP", 20)

    # split paragraphs, keep non-empty
    paragraphs = [p.strip() for p in re.split(r'\n{1,}', text) if p.strip()]
    sentences = []
    for p in paragraphs:
        # séparation naive en phrases — robuste pour la plupart des CV
        sents = re.split(r'(?<=[\.\?\!])\s+', p)
        sentences.extend([s.strip() for s in sents if s.strip()])

    # build chunks by accumulating sentences until approx chunk_size words
    chunks = []
    i = 0
    n = len(sentences)
    while i < n:
        current = []
        current_word_count = 0
        j = i
        while j < n:
            wcount = len(sentences[j].split())
            if current_word_count + wcount <= chunk_size or not current:
                current.append(sentences[j])
                current_word_count += wcount
                j += 1
            else:
                break
        chunks.append(" ".join(current))
        # compute next i with overlap (in words) -> keep tail overlap words
        if overlap > 0:
            tail_words = " ".join((" ".join(current)).split()[-overlap:])
            # start new chunk with tail_words as first sentence-like unit
            # to keep logic simple, we insert tail as pseudo-sentence then continue from j
            if j < n:
                sentences[j] = tail_words + " " + sentences[j]
            i = j
        else:
            i = j
    return chunks

def get_job_embedding(job: dict, sbert):
    """
    Retourne l'embedding du job (cache en mémoire pour éviter recalcule).
    La clé de cache est job.get('id') si présente, sinon un tuple title|company|description.
    """
    key = job.get("id")
    if not key:
        # safe key even for dicts: use title+company (if available)
        key = f"{job.get('title','')}_{job.get('company','')}_{len(job.get('description',''))}"
    if key in _job_emb_cache:
        return _job_emb_cache[key]
    job_txt = build_job_text(job)
    emb = sbert.encode(job_txt, convert_to_tensor=True)
    _job_emb_cache[key] = emb
    return emb

def score_semantic_enhanced(text: str, job: dict, sbert, top_k: int = 3) -> float:
    """
    Version améliorée de score_semantic:
    - accepte `job` (dict) et construit automatiquement job_txt
    - encode job une fois (cache)
    - chunking sémantique + sliding overlap
    - retourne une combinaison stable : 0.7*max + 0.3*mean(top_k)
    """
    try:
        if not text or not job:
            return 0.0
        job_emb = get_job_embedding(job, sbert)
        chunks = semantic_chunk_text(text)
        if not chunks:
            return 0.0
        sims = []
        # IMPORTANT: encode chunk par chunk (on encode en tenseur pour util.cos_sim)
        for chunk in chunks:
            chunk_emb = sbert.encode(chunk, convert_to_tensor=True)
            sims.append(util.cos_sim(chunk_emb, job_emb).item())
        sims_sorted = sorted(sims, reverse=True)
        max_sim = sims_sorted[0]
        mean_topk = sum(sims_sorted[:top_k]) / min(top_k, len(sims_sorted))
        combined = float(0.7 * max_sim + 0.3 * mean_topk)
        # clamp entre -1 et 1 (cosine) puis normaliser en 0..1 si tu préfères
        return max(-1.0, min(1.0, combined))
    except Exception as e:
        # en cas d'erreur, on degrade proprement
        print(f"[score_semantic_enhanced] erreur: {e}")
        return 0.0

def split_possible_lines(text):
    if not text:
        return []
    # prefer 'formation' block if present
    m = re.search(r"(?:formation[s]?\s*[:\-–])(.+)$", text, flags=re.I|re.S)
    t = m.group(1) if m else text
    parts = re.split(r"(?:\n|•|–|—|-|\d+\.)+", t)
    return [p.strip() for p in parts if p.strip()]

def canonical_degree_from_chunk(chunk):
    low = chunk.lower()
    for variant, canon in VARIANT_TO_CANON.items():
        if variant in low:
            return canon
    return None

def extract_years(s):
    return [int(y) for y in re.findall(r"(19|20)\d{2}", s)]

def guess_major_and_institution(chunk):
    parts = re.split(r"\(|,|:|-|—", chunk)
    parts = [p.strip() for p in parts if p.strip()]
    degree = parts[0] if parts else chunk
    major = None
    institution = None
    majors = ["data","science","informatique","finance","économie","business","management","marketing","statistique","math"]
    for p in parts[1:]:
        low = p.lower()
        for m in majors:
            if m in low:
                major = p
                break
        if major:
            break
    for p in parts[1:]:
        if any(k in p.lower() for k in ["univers","école","ecole","institut","college","facult","school","academy"]):
            institution = p
            break
    return degree, major, institution

def find_degree_terms(text):
    t = normalize_text(text)
    chunks = split_possible_lines(t)
    hits = []
    for chunk in chunks:
        if canonical_degree_from_chunk(chunk):
            hits.append(chunk)
    return hits

def field_match_score(major_string, job_fields):
    if not major_string:
        return 0.0
    s = normalize_text(major_string)
    job_fields_norm = [normalize_text(f) for f in (job_fields or [])]
    tokens = re.findall(r"\w+", s)
    score = 0.0
    for jf in job_fields_norm:
        for tok in tokens:
            if tok in jf or jf in tok:
                score = max(score, 1.0)
    if score == 0.0:
        matches = get_close_matches(s, job_fields_norm, n=1, cutoff=0.6)
        if matches:
            score = 0.8
    return score
# --- Robust education extraction helpers ---
def merge_edu_candidates(*lists):
    """
    Take several lists of tuples/dicts and return unified list of dicts:
    normalized keys: degree_canon, degree_raw, major, institution, years, chunk
    """
    out = []
    seen = set()
    for lst in lists:
        if not lst:
            continue
        for it in lst:
            # support tuple forms and dict forms
            if isinstance(it, dict):
                deg_canon = it.get("degree_canon")
                deg_raw = it.get("degree_raw") or it.get("degree") or ""
                major = it.get("major") or it.get("field") or ""
                inst = it.get("institution") or it.get("institution_name") or ""
                years = it.get("years") or it.get("year") or []
                chunk = it.get("chunk") or deg_raw or major or inst or ""
            elif isinstance(it, (list, tuple)):
                # try (deg, field, year)
                if len(it) >= 3:
                    deg_raw, major, years = it[0], it[1], it[2]
                elif len(it) == 2:
                    deg_raw, major = it[0], it[1]; years = []
                else:
                    deg_raw = it[0]; major = ""; years = []
                deg_canon = canonical_degree_from_chunk(deg_raw) or canonical_degree_from_chunk(major) or None
                inst = ""
                chunk = " | ".join([str(x) for x in (deg_raw, major) if x])
            else:
                continue

            key = (str(deg_canon or "").lower().strip(), str(deg_raw or "").lower().strip(), str(major or "").lower().strip(), str(inst or "").lower().strip())
            if key in seen:
                continue
            seen.add(key)
            out.append({
                "chunk": chunk,
                "degree_canon": deg_canon,
                "degree_raw": deg_raw,
                "major": major,
                "institution": inst,
                "years": years if isinstance(years, (list, tuple)) else ([years] if years else [])
            })
    return out


from datetime import datetime as _dt
import numpy as _np
import re
import unicodedata

def _make_jsonable(x):
    """Recursively convert common non-json types to json-serializable types."""
    if x is None:
        return None
    if isinstance(x, (str, bool, int, float)):
        return x
    if isinstance(x, (_np.integer,)):
        return int(x)
    if isinstance(x, (_np.floating,)):
        return float(x)
    if isinstance(x, (list, tuple, set)):
        return [_make_jsonable(i) for i in x]
    if isinstance(x, dict):
        return {str(k): _make_jsonable(v) for k, v in x.items()}
    try:
        if hasattr(x, "isoformat"):
            return x.isoformat()
    except Exception:
        pass
    return str(x)

def sanitize_section(section: str) -> str:
    """Normalize accents, collapse spaces, lowercase."""
    if not section:
        return ""
    text = unicodedata.normalize("NFKD", section)
    text = "".join(c for c in text if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", text).strip().lower()

def find_degree_field_pairs(text: str):
    """Find degree + preferred field combos in raw text."""
    DEGREE_TERMS = ["bachelor", "licence", "license", "master", "msc", "phd",
                        "doctorat", "dut", "bts", "deug", "baccalaureat"]
    results = []
    if not text:
        return results
    clean = sanitize_section(text)
    for deg in DEGREE_TERMS:
            # Match degree with possible 'en', 'in', 'dans', ':' before field
        pattern = (rf"{deg}"
          r"(?:\s+de\s+(?:recherche|professionnelle|specialise[e]?)\s*)?"
          r"\s*(en|in|dans|:)\s+([a-zA-Z0-9\-\s]+)"
        )


        for m in re.finditer(pattern, clean):
            degree_raw = deg
            major = m.group(2).strip()
            results.append({
                "degree_raw": degree_raw,
                "degree_canon": canonical_degree_from_chunk(degree_raw),
                "major": major,
                "institution": "",
                "years": []
            })
    return results
# --------------------- MAIN ---------------------
def main(resume_dir: str, job_input, file_list=None):
    import os, json, logging, re
    from datetime import datetime as _dt
    import numpy as _np

    # helper: json-able conversion
    def _make_jsonable(x):
        if x is None:
            return None
        if isinstance(x, (str, bool, int, float)):
            return x
        if isinstance(x, (_np.integer,)):
            return int(x)
        if isinstance(x, (_np.floating,)):
            return float(x)
        if isinstance(x, (list, tuple, set)):
            return [_make_jsonable(i) for i in x]
        if isinstance(x, dict):
            return {str(k): _make_jsonable(v) for k, v in x.items()}
        try:
            if hasattr(x, "isoformat"):
                return x.isoformat()
        except Exception:
            pass
        return str(x)

    # defensive logger
    log = logging.getLogger("ResumeMatcher.main")

    # seed / models
    try:
        DetectorFactory.seed = 0
    except Exception:
        log.debug("DetectorFactory not available or seed failed.", exc_info=True)

    # load sbert & spacy defensively
        # load sbert & spacy defensively
    DEFAULT_SBERT = "sentence-transformers/distiluse-base-multilingual-cased-v2"
    model_name = getattr(Config, "MODEL_NAME", DEFAULT_SBERT)
    try:
        sbert = SentenceTransformer(model_name)
    except Exception as e:
        # utiliser la variable locale model_name (toujours présente) pour le log
        log.warning("Failed to load SBERT (%s). Continuing with sbert=None. Error: %s", model_name, e)
        sbert = None

    try:
        nlp = spacy.load("fr_core_news_sm")
    except Exception:
        log.warning("Failed to load spacy fr_core_news_sm, continuing without nlp", exc_info=True)
        nlp = None

    try:
        nlp = spacy.load("fr_core_news_sm")
    except Exception:
        log.warning("Failed to load spacy fr_core_news_sm, continuing without nlp", exc_info=True)
        nlp = None

    # load job config
    with open(job_input, "r", encoding="utf-8") as f:
        job = json.load(f)

    summary_rows = []
    detail_rows = []

    # determine resume list (prefer provided file_list)
    if file_list:
        resume_files = [f for f in file_list if f.lower().endswith((".pdf", ".docx", ".txt", ".jpg", ".jpeg", ".png"))]
    else:
        resume_files = [f for f in os.listdir(resume_dir) if f.lower().endswith((".pdf", ".docx", ".txt", ".jpg", ".jpeg", ".png"))]

    if not resume_files:
        log.warning("No resume files found in %s (file_list=%s)", resume_dir, bool(file_list))
        return pd.DataFrame(), pd.DataFrame()

    for fn in resume_files:
        path = os.path.join(resume_dir, fn)
        log.info("Start processing file %s", fn)

        try:
            raw_text = extract_text_from_file(path)
            cleaned = nettoyer_texte(raw_text)


    # If you want translation as well:
    # rawo = extract_and_translate(path) or ""
    # rawz = normalize_text(rawo)

    # Clean your text (your existing cleaning function)
            print("raw used for education scoring")
            print(repr(raw_text))   # repr() shows full string including line breaks
            print("the cleaned")
            print(repr(cleaned))

            txt = cleaned
        except Exception as e:
            log.exception("process_file failed for %s: %s", fn, e)
            txt = ""

        # classification
        try:
          predicted_class, class_top_prob, class_prob_dict = classify_cv(txt)
          #predicted_class, class_top_prob, class_prob_dict = classify_cv(txt, is_french=True)
        except Exception:
            log.exception("classify_cv failed for %s", fn)
            predicted_class, class_top_prob, class_prob_dict = "unknown", 0.0, {}

        # summary (kept for UI/semantic but not used for edu strict matching)
        try:
            rsummary = generate_summary(txt)
            resume_summary=clean_summary(rsummary)
        except Exception as e:
            log.exception("generate_summary failed for %s: %s", fn, e)
            resume_summary = ""

        # languages & skills
        try:
            lang_sc, langs = score_language(txt or "", set(job.get("required_languages", [])))
        except Exception:
            log.exception("score_language failed for %s", fn)
            lang_sc, langs = 0.0, set()

        try:
            skill_sc, skills = score_skills(txt or "", job.get("required_skills", []))
            skills = list(skills or [])
        except Exception:
            log.exception("score_skills failed for %s", fn)
            skill_sc, skills = 0.0, []
#mnhna
# --- EDUCATION extraction & scoring (monolithic) ---
        try:
    # Normalize job inputs
          job_required_degree = job.get("required_degree", "").strip().lower()
          preferred_fields = [f.strip().lower() for f in job.get("preferred_fields", [])]

    # Call the new education scoring
          edu_result = score_education(
            resume_text=raw_text,
            job_required_degree=job_required_degree,
            preferred_fields=preferred_fields
          )

    # Debug prints (optional)
          print("Resume Text:", txt[:], "...")  # show first 300 chars
          print("Required degree:", job_required_degree)
          print("Preferred fields:", preferred_fields)

    # Extract score + best match
          edu_sc = float(edu_result.get("score", 0.0))
          best_hit = edu_result.get("best")  # already normalized by function

    # Build matched_education_list
          matched_education_list = []
          if best_hit:
              matched_education_list.append({
                "degree": best_hit.get("degree_canon") or best_hit.get("degree_raw") or "",
                "field": best_hit.get("degree_raw") or "",
                "year": best_hit.get("years") or []
            })

        except Exception as e:
           log.exception("Education analysis failed for %s: %s", fn, e)
           edu_sc = 0.0
           matched_education_list = []

# Ensure JSON-serializable
        try:
           matched_education_list = _make_jsonable(matched_education_list)
        except Exception:
           log.exception("Make jsonable failed for matched_education_list")

# Ensure JSON-serializable for downstream usage
        """
        try:
           matched_education_list = _make_jsonable(matched_education_list)
        except Exception:
           log.exception("Make jsonable failed for matched_education_list")
        """
        # semantic
        try:
            sem_sc = score_semantic_enhanced(txt or "", job, sbert)

        except Exception:
            log.exception("score_semantic failed for %s", fn)
            sem_sc = 0.0
        
        
        #exper
        # ---------- EXPERIENCE (full_experience) ----------
        try:
            exp_result = full_experience(
               raw_text=txt,  # the normalized raw text you already extracted above
               min_years=job.get("min_experience_years", 0),
               max_years=job.get("max_experience_years", 50),
               job_titles=job.get("experience_fields", []) or job.get("description", "")
            )

            exp_sc = exp_result.get("score", 0.0)
            matched_exps_for_output = exp_result.get("entries", [])
            total_relevant_years = exp_result.get("total_years", 0.0)

        except Exception as e:
           log.exception("Experience analysis failed for %s: %s", fn, e)
           exp_sc = 0.0
           matched_exps_for_output = []
           total_relevant_years = 0.0

        # inside your main resume loop, after edu_sc, sem_sc, skill_sc, etc.
        LLAMA_ENABLED = True  # set to False to skip LLaMA scoring
        llama_sc = 0.0
        llama_expl = ""

        if LLAMA_ENABLED:
            try:
                llama_sc, llama_expl = llama_score_resume(
                   resume_text=txt or "",
                   job=job,
                   api_keys=[os.environ.get("LLM_API_KEY")] if os.environ.get("LLM_API_KEY") else None,
                   model_id="llama-3.3-70b-versatile"
                )
            except Exception as e:
              log.exception("llama_score_resume failed for %s: %s", fn, e)
              llama_sc = 0.0
              llama_expl = f"Error: {e}"

# --- base score (without LLaMA)
        base_weights = ["semantic", "skills", "experience", "education", "language"]
        weights_total = sum([Config.WEIGHTS.get(k, 0) for k in base_weights])
        base_score = (
               Config.WEIGHTS.get("semantic", 0) * float(sem_sc) +
               Config.WEIGHTS.get("skills", 0) * float(skill_sc) +
               Config.WEIGHTS.get("experience", 0) * float(exp_sc) +
               Config.WEIGHTS.get("education", 0) * float(edu_sc) +
               Config.WEIGHTS.get("language", 0) * float(lang_sc)
            ) / max(weights_total, 1e-6)
        """
# --- smart LLaMA blending
        if LLAMA_ENABLED:
    # Normalize LLaMA score (if needed)
            llama_score_norm = float(llama_sc)
    
    # Dynamic weight: higher when base_score is low or LLaMA is extreme
            base_confidence = base_score  # 0-1 scale expected
            dynamic_weight = 0.5 + 0.5 * abs(llama_score_norm - 0.5)  # ranges 0.5-1
            max_llama_weight = Config.WEIGHTS.get("llama", 0.9)  # max influence
            llama_weight = min(dynamic_weight, max_llama_weight)
    
    # Blend scores
            final_score = round((1 - llama_weight) * base_score + llama_weight * llama_score_norm, 4)
        else:
           final_score = round(base_score, 4)
        """
        if LLAMA_ENABLED:
            llama_score_norm = float(llama_sc)
    
    # Base dynamic weight: extreme LLaMA scores get more weight
            dynamic_weight = 0.5 + 0.5 * abs(llama_score_norm - 0.5)
    
    # Increase weight if base_score is low (parser unsure)
            dynamic_weight += (1 - base_score) * 0.2
    
            max_llama_weight = Config.WEIGHTS.get("llama", 0.85)
            llama_weight = min(dynamic_weight, max_llama_weight)
    
    # Threshold override for very confident LLaMA
            if llama_score_norm > 0.85 or llama_score_norm < 0.15:
                final_score = round(llama_score_norm, 4)
            else:
                final_score = round((1 - llama_weight) * base_score + llama_weight * llama_score_norm, 4)
        else:
            final_score = round(base_score, 4)

# --- append to summary
        summary_rows.append({
          "file": fn,
          "summary": resume_summary,
          "semantic": float(sem_sc),
          "skills": float(skill_sc),
          "experience": float(exp_sc),
          "education": float(edu_sc),
          "language": float(lang_sc),
          "llama": float(llama_sc or 0.0),
          "llama_expl": llama_expl,
          "predicted_class": predicted_class,
          "class_top_prob": round(class_top_prob, 4),
          "final": float(final_score)
       })

# --- detail rows (mostly unchanged)
        detail_rows.append({
          "file": fn,
          "languages": _make_jsonable(list(langs or [])),
          "languages_str": ",".join(langs) if langs else "",
          "matched_skills": ",".join(sorted(set(skills))) if skills else "",
          "matched_skills_list": _make_jsonable(list(skills or [])),
          "class_probs": _make_jsonable(class_prob_dict or {}),
          "matched_education": "; ".join([f"{e.get('degree','')}:{e.get('field','')}" for e in (matched_education_list or [])]),
          "matched_education_list": matched_education_list,
          "matched_experience_domains": matched_exps_for_output,
          "total_relevant_years": total_relevant_years,
          "llama_expl": llama_expl
       })

# --- final dataframes
        df_sum = pd.DataFrame(summary_rows).sort_values("final", ascending=False, ignore_index=True) if summary_rows else pd.DataFrame()
        df_det = pd.DataFrame(detail_rows).set_index("file") if detail_rows else pd.DataFrame()
        return df_sum, df_det

